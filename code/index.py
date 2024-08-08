from tkinter import *
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename
from docx import Document
import re
import os

window = Tk()

class Formater:
    def cpf_formater(text, var, index, mode): #TODO cpf formater
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 11:
           valor = valor[:3] + "." + valor[3:6] + "." + valor[6:9] + "-" + valor[9:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)

    def cnpj_formater(text, var, index, mode): #TODO cpf formater
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 11:
           valor = valor[:3] + "." + valor[3:6] + "." + valor[6:9] + "-" + valor[9:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)

    def cep_formater(text, var, index, mode): #TODO cpf formater
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 11:
           valor = valor[:3] + "." + valor[3:6] + "." + valor[6:9] + "-" + valor[9:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)

    def rg_formater(text, var, index, mode): #TODO cpf formater
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 11:
           valor = valor[:3] + "." + valor[3:6] + "." + valor[6:9] + "-" + valor[9:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)
    
    def date_formater(text, var, index, mode): #TODO cpf formater
        #Só recebe valor que passa pelo validador
        valor = text.get()
        if len(valor) == 11:
           valor = valor[:3] + "." + valor[3:6] + "." + valor[6:9] + "-" + valor[9:]
        else:
            valor = valor.replace('.','').replace('-','')
        text.set(valor)
        

class Validator:
    def str_validator(text):
        return not text.isdecimal()
    
    def num_validator(text):
        return text.isdecimal()
    
    def cpf_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 15:
            if len(text) >= 12:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
        
    def cnpj_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 15:
            if len(text) >= 12:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
    
    def cep_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 15:
            if len(text) >= 12:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
    
    def rg_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 15:
            if len(text) >= 12:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False
    
    def date_validator(text):
        padrao = r"^[-\d.,/]+$"  # Permite dígitos, ponto, vírgula, hífen e barra
        if len(text) < 15:
            if len(text) >= 12:
                return re.match(padrao, text) is not None
            else:
                return text.isdecimal()
        return False

class App:
    def __init__(self):
        self.window = window
        self.tela()
        self.pageMenu()
        window.mainloop()

    def tela(self):
        self.window.configure(background='darkblue')
        self.window.resizable(False,False)
        self.window.geometry('880x500')
        self.window.iconbitmap('./code/imgs/delta-icon.ico')
        self.window.title('Gerador de CPS')

    def pageMenu(self):
        self.menu = Frame(self.window, bd=4, bg='lightblue')
        self.menu.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        self.textOrientacao = Label(self.menu, text='Selecione o tipo de CPS que deseja fazer:', background='lightblue', font=('arial',20,'bold'))\
        .place(relx=0.15,rely=0.2,relheight=0.15)
        
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice-hori.png')
        
        self.logo = self.logo.subsample(4,4)
        
        Label(self.window, image=self.logo, background='lightblue')\
            .place(relx=0.175,rely=0.05,relwidth=0.7,relheight=0.2)

        #Pessoa física
        self.btnPF = Button(self.menu, text='CPS Pessoa Física',\
            command= lambda: self.pagePF())\
                .place(relx=0.15,rely=0.7,relwidth=0.25,relheight=0.15)

        #Inatividade
        self.btnIN = Button(self.menu, text='CPS Inatividade',\
            command= lambda: self.pageIN())\
                .place(relx=0.60,rely=0.4,relwidth=0.25,relheight=0.15)

        #Lucro Presumido
        self.btnLP = Button(self.menu, text='CPS Lucro Presumido',\
            command= lambda: self.pageLP())\
                .place(relx=0.15,rely=0.4,relwidth=0.25,relheight=0.15)

        #Simples Nacional
        self.btnSN = Button(self.menu, text='CPS Simples Nacional',\
            command= lambda: self.pageSN())\
                .place(relx=0.60,rely=0.7,relwidth=0.25,relheight=0.15)

    def pagePF(self):
        self.menu.destroy()
        self.doc = Document('./code/CPS\'s/CPS PESSOA FISICA.docx')

        self.cpsPF = Frame(self.window, bd=4, bg='lightblue')
        self.cpsPF.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        #Titulo
        Label(self.cpsPF, text='Gerador de CPS Pessoa física', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.05)

        #Botão voltar
        Button(self.cpsPF, text='Voltar ao menu',\
            command= lambda: (self.cpsPF.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeContra' : StringVar(), #strValidator
            '$estadoContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            "$compleContra" : StringVar(), #strValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$sspContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            "$numEmpre" : StringVar(),  #dateValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Contratante
        Label(self.cpsPF, text='Contratante',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.2)

        ###########nome

        Label(self.cpsPF, text='Nome',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.3)

        self.nomeEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$nomeContra'])\
                .place(relx=0.05,rely=0.37,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsPF, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.30)

        self.ruaEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$ruaContra'])\
                .place(relx=0.35,rely=0.37,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsPF, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.30)

        self.numEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$numContra'])\
                .place(relx=0.61,rely=0.37,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsPF, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.30)

        self.bairroEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$bairroContra'])\
                .place(relx=0.7,rely=0.37,relwidth=0.25,relheight=0.05)

        ###########CEP

        Label(self.cpsPF, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.45)

        self.cepEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$cepContra'])\
                .place(relx=0.05,rely=0.52,relwidth=0.25,relheight=0.05)
        
        ###########Estado Civil

        Label(self.cpsPF, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.45)

        self.estadoEntry = StringVar(self.cpsPF)

        self.estadoEntryOpt = ('solteiro(a)','divorsiado(a)','viuvo(a)')

        self.estadoEntry.set('')

        self.popup = OptionMenu(self.cpsPF, self.estadoEntry, *self.estadoEntryOpt)
        
        self.menuCasado = self.popup['menu']

        #Casado
        self.subLista = Menu(self.menuCasado, tearoff=False)
        self.menuCasado.add_cascade(label = 'casado(a)',menu= self.subLista)
        self.subLista.add_command(label='Comunhão Parcial de Bens', \
            command= lambda: self.estadoEntry.set('casado(a) em CPB'))
        
        self.subLista.add_command(label='Comunhão Total de Bens',\
            command= lambda: self.estadoEntry.set('casado(a) em CTB'))
        
        self.subLista.add_command(label='Separação Total de Bens',\
            command= lambda: self.estadoEntry.set('casado(a) em STB'))


        self.popup.place(relx=0.35,rely=0.52,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoContra'] = self.estadoEntry

        ###########Complemento

        Label(self.cpsPF, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.45)

        self.complementoEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.52,relwidth=0.34,relheight=0.05)

        ###########RG

        Label(self.cpsPF, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.60)

        self.rgEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$rgContra'])\
                .place(relx=0.05,rely=0.67,relwidth=0.20,relheight=0.05)
        
        ###########Num. Empregados

        Label(self.cpsPF, text='Num. Empregados',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.6)

        self.numEmpreEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$numEmpre'])\
                .place(relx=0.35,rely=0.67,relwidth=0.05,relheight=0.05)

        
        ###########CPF

        Label(self.cpsPF, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.61,rely=0.6)

        self.cpfEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$cpfContra'])\
                .place(relx=0.61,rely=0.67,relwidth=0.25,relheight=0.05)
        
        
        ###########Org. Emissor

        Label(self.cpsPF, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.9,rely=0.6)

        self.sspEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$sspContra'])\
                .place(relx=0.9,rely=0.67,relwidth=0.05,relheight=0.05)
        

        #Contrato
        Label(self.cpsPF, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.75)

        ###########Valor pagamento

        Label(self.cpsPF, text='Val. pagamento',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.85)

        self.valPagEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$valPag'])\
                .place(relx=0.05,rely=0.92,relwidth=0.15,relheight=0.05)

        ###########Data inicio

        Label(self.cpsPF, text='Dia início',\
            background='lightblue', font=(10))\
                .place(relx=0.25,rely=0.85)

        self.dtInicEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$dtInic'])\
                .place(relx=0.25,rely=0.92,relwidth=0.1,relheight=0.05)

        ###########Data pagamento

        Label(self.cpsPF, text='Dia vencimento',\
            background='lightblue', font=(10))\
                .place(relx=0.4,rely=0.85)

        self.dtPagEntry = Entry(self.cpsPF,\
            textvariable=self.referencias['$dtVenc'])\
                .place(relx=0.4,rely=0.92,relwidth=0.1,relheight=0.05)

        #Botão enviar
        Button(self.cpsPF, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.cpsPF))\
                .place(relx=0.61,rely=0.85,relwidth=0.35,relheight=0.12)

    def pageIN(self):
        self.menu.destroy()
        self.doc = Document('./code/CPS\'s/CPS INATIVIDADE.docx')

        self.cpsIN = Frame(self.window, bd=4, bg='lightblue')
        self.cpsIN.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        #Titulo
        Label(self.cpsIN, text='Gerador de CPS Inatividade', background='lightblue', font=('arial',17,'bold'))\
            .place(relx=0.3,rely=0.045)
            
        #Logo
        self.logo = PhotoImage(file='./code/imgs/deltaprice_logo-slim.png')
        
        self.logo = self.logo.subsample(5,5)
        
        Label(self.cpsIN, image=self.logo, background='lightblue')\
            .place(relx=0.75,rely=0.01,relwidth=0.12,relheight=0.15)

        #Botão voltar
        Button(self.cpsIN, text='Voltar ao menu',\
            command= lambda: (self.cpsIN.destroy, self.pageMenu()))\
                .place(relx=0,rely=0,relwidth=0.25,relheight=0.06)

        self.referencias = {
            '$nomeEmp' : StringVar(), #strValidator
            '$ruaEmp' : StringVar(), #strValidator
            '$numEmp' : StringVar(), #intValidator
            '$bairroEmp' : StringVar(),  #strValidator
            '$cepEmp' : StringVar(),  #intValidator
            '$rgEmp' : StringVar(),  #rgValidator
            '$sspEmp' : StringVar(),  #sspValidator
            '$cnpjEmp' : StringVar(),  #cpfValidator
            '$nomeContra' : StringVar(), #strValidator
            '$ruaContra' : StringVar(), #strValidator
            '$numContra' : StringVar(), #intValidator
            '$bairroContra' : StringVar(),  #strValidator
            '$cepContra' : StringVar(),  #intValidator
            '$rgContra' : StringVar(),  #rgValidator
            '$sspContra' : StringVar(),  #sspValidator
            '$cpfContra' : StringVar(),  #cpfValidator
            '$estadoContra' : StringVar(), #strValidator
            "$compleEmp" : StringVar(), #strValidator
            "$compleContra" : StringVar(), #strValidator
            "$dtVenc" : StringVar(),  #dateValidator
            "$valPag" : StringVar(), #intValidator
            "$dtInic" : StringVar()  #dateValidator
        }

        #Labels e Entrys
        #Empresa
        Label(self.cpsIN, text='Empresa',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.115)
                
        self.canvas = Canvas(self.cpsIN, width=625, height=10, background='darkblue',border=-5)
        self.canvas.place(relx=0.17,rely=0.15)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y
                
        ###########nome empresa

        Label(self.cpsIN, text='Nome empresa',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.18)

        self.nomeEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$nomeEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.23,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsIN, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.18)

        self.ruaEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$ruaEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.23,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsIN, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.18)

        self.numEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$numEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                    .place(relx=0.61,rely=0.23,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsIN, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.18)

        self.bairroEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$bairroEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.23,relwidth=0.25,relheight=0.05)

        ###########CEP

        Label(self.cpsIN, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.31)

        self.cepEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$cepEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.36,relwidth=0.25,relheight=0.05)

        ###########CNPJ

        Label(self.cpsIN, text='CNPJ',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.31)

        self.cnpjEmpEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$cnpjEmp'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.36,relwidth=0.2,relheight=0.05)
                
        ###########Complemento

        Label(self.cpsIN, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.31)

        self.complementoEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$compleEmp'])\
                .place(relx=0.61,rely=0.36,relwidth=0.35,relheight=0.05)
        
        #Socio
        Label(self.cpsIN, text='Sócio',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.42)
                
        self.canvas = Canvas(self.cpsIN, width=655, height=10,border=-5)
        self.canvas.place(relx=0.13,rely=0.455)
                
        self.canvas.create_line(-5,0,655,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########nome

        Label(self.cpsIN, text='Nome sócio',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.48)

        self.nomeEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$nomeContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########rua

        Label(self.cpsIN, text='Rua',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.48)

        self.ruaEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$ruaContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.53,relwidth=0.20,relheight=0.05)

        ###########Num

        Label(self.cpsIN, text='Num.',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.48)

        self.numEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$numContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.61,rely=0.53,relwidth=0.05,relheight=0.05)

        ###########bairro

        Label(self.cpsIN, text='Bairro',\
            background='lightblue', font=(10))\
                .place(relx=0.7,rely=0.48)

        self.bairroEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$bairroContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.7,rely=0.53,relwidth=0.25,relheight=0.05)

        ###########CEP

        Label(self.cpsIN, text='CEP',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.61)

        self.cepEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$cepContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.66,relwidth=0.25,relheight=0.05)

        ###########RG

        Label(self.cpsIN, text='RG',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.61)

        self.rgEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$rgContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.35,rely=0.66,relwidth=0.20,relheight=0.05)

        ###########Org. Emissor

        Label(self.cpsIN, text='Org.',\
            background='lightblue', font=(10))\
                .place(relx=0.9,rely=0.61)

        self.sspEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$sspContra'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: not text.isdecimal()), '%S'))\
                .place(relx=0.9,rely=0.66,relwidth=0.05,relheight=0.05)

        ###########TODO CPF

        self.val = StringVar()

        self.val.trace_add('write', lambda *args, passed = self.val:\
            Formater.cpf_formater(passed, *args) )

        Label(self.cpsIN, text='CPF',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.61)
        

        self.cpfEntry = Entry(self.cpsIN, textvariable = self.val, \
            validate ='key', validatecommand =(self.cpsIN.register(Validator.cpf_validator), '%P'))\
                .place(relx=0.61,rely=0.66,relwidth=0.25,relheight=0.05)

        self.referencias['$cpfContra'] = self.val

        ###########Estado Civil

        Label(self.cpsIN, text='Estado Civil',\
            background='lightblue', font=(10))\
                .place(relx=0.35,rely=0.72)

        self.estadoEntry = StringVar(self.cpsIN)

        self.estadoEntryOpt = ('solteiro(a)', 'casado(a)','divorsiado(a)','viuvo(a)')

        self.estadoEntry.set('solteiro(a)')

        self.popup = OptionMenu(self.cpsIN, self.estadoEntry, *self.estadoEntryOpt)\
            .place(relx=0.35,rely=0.77,relwidth=0.2,relheight=0.06)

        self.referencias['$estadoContra'] = self.estadoEntry

        ###########Complemento

        Label(self.cpsIN, text='Complemento',\
            background='lightblue', font=(10))\
                .place(relx=0.6,rely=0.73)

        self.complementoEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$compleContra'])\
                .place(relx=0.61,rely=0.78,relwidth=0.35,relheight=0.05)

        #Contrato
        Label(self.cpsIN, text='Contrato',\
            background='lightblue', font=('Times New Roman',15,'bold italic'))\
                .place(relx=0.05,rely=0.81)
                
        self.canvas = Canvas(self.cpsIN, width=625, height=10,border=-5)
        self.canvas.place(relx=0.17,rely=0.845)
                
        self.canvas.create_line(-5,0,625,0, fill="darkblue", width=10)
        
        # x,angulo x , y, angulo y

        ###########Valor pagamento

        Label(self.cpsIN, text='Val. pagamento',\
            background='lightblue', font=(10))\
                .place(relx=0.05,rely=0.88)

        self.valPagEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$valPag'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.05,rely=0.93,relwidth=0.15,relheight=0.05)

        ###########Data inicio

        Label(self.cpsIN, text='Dia início',\
            background='lightblue', font=(10))\
                .place(relx=0.25,rely=0.88)

        self.dtInicEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$dtInic'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.25,rely=0.93,relwidth=0.1,relheight=0.05)

        ###########Data pagamento

        Label(self.cpsIN, text='Dia vencimento',\
            background='lightblue', font=(10))\
                .place(relx=0.4,rely=0.88)

        self.dtPagEntry = Entry(self.cpsIN,\
            textvariable=self.referencias['$dtVenc'],\
                validate='key', validatecommand=(self.cpsIN.register(lambda text: text.isdecimal()), '%S'))\
                .place(relx=0.4,rely=0.93,relwidth=0.1,relheight=0.05)

        #Botão enviar
        Button(self.cpsIN, text='Gerar CPS',\
            command= lambda: self.alterar_doc(frame_ativo=self.cpsIN))\
                .place(relx=0.61,rely=0.865,relwidth=0.35,relheight=0.12)

    def pageLP(self):
        self.menu.destroy()

        self.cpsLP = Frame(self.window, bd=4, bg='lightblue')
        self.cpsLP.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        self.textOrientacao = Label(self.cpsLP, text='OI', background='lightblue', font=('Bold', 15))\
            .place(relx=0.22,rely=0.1,relheight=0.15)

        #Botão voltar
        self.btnVoltar = Button(self.cpsLP, text='Voltar ao menu',\
            command= lambda: (self.cpsLP.destroy, self.pageMenu()))\
                .place(relx=0.15,rely=0.7,relwidth=0.25,relheight=0.15)

    def pageSN(self):
        self.menu.destroy()

        self.cpsSN = Frame(self.window, bd=4, bg='lightblue')
        self.cpsSN.place(relx=0.05,rely=0.05,relwidth=0.9,relheight=0.9)

        self.textOrientacao = Label(self.cpsSN, text='OI', background='lightblue', font=('Bold', 15))\
            .place(relx=0.22,rely=0.1,relheight=0.15)

        #Botão voltar
        self.btnVoltar = Button(self.cpsSN, text='Voltar ao menu',\
            command= lambda: (self.cpsSN.destroy, self.pageMenu()))\
                .place(relx=0.15,rely=0.7,relwidth=0.25,relheight=0.15)

    def alterar_doc(self, frame_ativo):
        for par in self.doc.paragraphs:
            for itens in self.referencias:
                if par.text.find(itens) != -1:
                    par.text = par.text.replace(itens, self.referencias[itens].get())


        file = asksaveasfilename(title='Defina o nome e o local onde o arquivo será salvo', filetypes=((".docx","*.docx"),))

        self.doc.save(file+'.docx')

        os.startfile(file+'.docx')

        frame_ativo.destroy()
        self.pageMenu()

App()